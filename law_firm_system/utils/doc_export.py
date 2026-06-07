"""
文书导出工具：支持导出Word(docx)和PDF格式
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

def export_docx(template_content, case_info, firm_info, output_path):
    """
    导出Word文档
    :param template_content: 模板内容(HTML格式，需转换)
    :param case_info: 案件信息字典
    :param firm_info: 律所信息字典
    :param output_path: 输出文件路径
    :return: 成功返回True，失败返回False
    """
    try:
        doc = Document()
        
        # 设置文档默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        
        # 添加标题
        heading = doc.add_heading('法律文书', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 填充模板变量
        content = template_content
        if case_info:
            content = content.replace('{{客户名称}}', case_info.get('client_name', ''))
            content = content.replace('{{案号}}', case_info.get('case_no', ''))
            content = content.replace('{{受理法院}}', case_info.get('court', ''))
            content = content.replace('{{律所名称}}', firm_info.get('name', '') if firm_info else '')
            content = content.replace('{{承办律师}}', case_info.get('lawyer_name', ''))
        
        # 添加正文（简单处理，实际应解析HTML）
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para)
        
        # 保存文档
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"导出Word失败: {e}")
        return False

def export_pdf(html_content, output_path):
    """
    导出PDF文档（需要额外依赖如weasyprint或pdfkit）
    这里先预留接口
    """
    # TODO: 实现PDF导出功能
    pass
